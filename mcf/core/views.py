from dataclasses import fields
from msilib.schema import SelfReg
from multiprocessing import Value
from typing import Self
from django.shortcuts import render
from django.shortcuts import render, redirect
from docx import Document
from .forms import Client_Form_Info, SearchForm
from .models import Client_Account, Client_Form_Info, Excel_Report
from django.http import HttpResponse, FileResponse
import os
from django.urls import reverse
from django.shortcuts import redirect
from django import forms
from django.db import IntegrityError
from django.db import models
from django.contrib import messages
from io import BytesIO
from openpyxl import Workbook
from django.http import JsonResponse
from django.shortcuts import render, get_object_or_404
from django.core import serializers
from django.views.generic import View 
from django.template import loader
from django.db.models import Q, Count, Case, When
from django.core.exceptions import ObjectDoesNotExist
from django.shortcuts import render


def index(request):
    # Your view logic here
    return render(request, 'core/index.html')


from django.db.models import Count, Sum

from django.contrib.humanize.templatetags.humanize import intcomma

def dashboard(request):
    # Pie chart for gender
    gender_data = list(Client_Account.objects.values('GENDER').annotate(count=Count('GENDER')))
    # Bar graph for each SOL_ID
    sol_id_data = list(Client_Account.objects.values('SOL_ID').annotate(count=Count('SOL_ID')))
    # Filtering sum of loans worked on
    filtered_accounts_acctnum = Client_Form_Info.objects.values_list('ACCTNUM', flat=True)
    filtered_accounts = Client_Account.objects.filter(ACCTNUM__in=filtered_accounts_acctnum)
    total_claimed_amount = filtered_accounts.aggregate(total_claimed=Sum('AMOUNT_CLAIMED'))['total_claimed']
    if total_claimed_amount:
        total_claimed_amount = intcomma(int(total_claimed_amount))  # Add thousand separators
    else:
        total_claimed_amount = 0
    total_claimed_count = filtered_accounts.count()
    # Line graph for age categories (assuming you have a function to categorize age)
    age_category_data = list(Client_Account.objects.values('AGE_CATEGORY').annotate(count=Count('AGE_CATEGORY')))
    context = {
        'gender_data': gender_data,
        'sol_id_data': sol_id_data,
        'total_claimed_amount': total_claimed_amount,
        'total_claimed_count': total_claimed_count,
        'age_category_data': age_category_data,
    }
    return render(request, 'core/dashboard.html', context)
from django.shortcuts import render, get_object_or_404
from django.core.exceptions import ObjectDoesNotExist
from .models import Client_Account, Client_Form_Info, Excel_Report

def full_details(request, account_num):
    client_account = get_object_or_404(Client_Account, ACCTNUM=account_num)

    try:
        client_form_info = Client_Form_Info.objects.get(ACCTNUM=account_num)
    except ObjectDoesNotExist:
        client_form_info = None


    client_details = [
        # Client_Account model fields
        {'name': 'SOL ID', 'value': client_account.SOL_ID or ''},
        {'name': 'CIF ID', 'value': client_account.CIF_ID or ''},
        {'name': 'Account Manager User ID', 'value': client_account.ACCT_MGR_USER_ID or ''},
        {'name': 'Account Number', 'value': client_account.ACCTNUM or ''},
        {'name': 'Account Name', 'value': client_account.ACCT_NAME or ''},
        {'name': 'Arrears Days', 'value': client_account.ARREARSDAYS or ''},
        # Add more fields as needed

        # Client_Form_Info model fields
        {'name': 'Loan Application Date', 'value': client_form_info.LOAN_APP_DATE or '' if client_form_info else ''},
        {'name': 'Address', 'value': client_form_info.ADDRESS or '' if client_form_info else ''},
        {'name': 'Loan Purpose', 'value': client_form_info.LOAN_PURPOSE or '' if client_form_info else ''},
        {'name': 'Group', 'value': client_form_info.GROUP or '' if client_form_info else ''},
        {'name': 'Reason for Default', 'value': client_form_info.REASON_FOR_DEFAULT or '' if client_form_info else ''},
        {'name': 'Contact', 'value': client_form_info.CONTACT or '' if client_form_info else ''},

    ]

    return render(request, 'core/full_details.html', {'client_details': client_details})
from django.db.models import Case, When, Value, IntegerField, CharField, F
from django.contrib.humanize.templatetags.humanize import intcomma

def search_client_details(request):
    if request.method == 'GET':
        search_query = request.GET.get('q', '')
        if search_query:
            clients = Client_Account.objects.filter(
                Q(ACCT_NAME__icontains=search_query) |
                Q(ACCTNUM__icontains=search_query) |
                Q(CIF_ID__icontains=search_query) |
                Q(NATIONAL_ID_REG_NUMBER__icontains=search_query) |
                Q(AMOUNT_CLAIMED__icontains=search_query)
            )
            """if Client_Account.AMOUNT_CLAIMED:
                Client_Account.AMOUNT_CLAIMED = intcomma(int(Client_Account.AMOUNT_CLAIMED))
            else:
                Client_Account.AMOUNT_CLAIMED = 0"""
            STATUS=Case(
                    When(ACCTNUM__in=Client_Form_Info.objects.values_list('ACCTNUM', flat=True), then=Value('Done')),
                    default=Value('Pending'),
                    output_field=forms.CharField(),
                )
            return render(request, 'core/search_results.html', {'clients': clients, 'search_query': search_query})
        else:
            return render(request, 'core/search_form.html')
    else:
        return render(request, 'core/search_form.html')

def search_client(request):
    if request.method == 'POST':
        form = SearchForm(request.POST)
        if form.is_valid():
            account_number = form.cleaned_data['account_number']
            try:
                client_account = Client_Account.objects.get(ACCTNUM=account_number)  # Removed unnecessary ACCTNUM__isnull=False

                # Redirect to the add_client_info view, passing the account_number
                return redirect(reverse('add_client_info', kwargs={'account_number': account_number}))

            except Client_Account.DoesNotExist:
                error_message = 'Client account not found.'
                return render(request, 'core/search_client.html', {'form': form, 'error_message': error_message})

    else:
        form = SearchForm()

    return render(request, 'core/search_client.html', {'form': form})






def add_client_info(request, account_number):
    context = {}

    try:
        client_account = Client_Account.objects.get(ACCTNUM=account_number)
        # Create form based on Client_Form_Info with fields
        Client_Info_Form = forms.modelform_factory(Client_Form_Info, fields='__all__')
        form = Client_Info_Form(initial={'ACCTNUM': Client_Account.objects.get(ACCTNUM=account_number)})
        if request.method == 'POST':
                form = Client_Info_Form(request.POST)
                if form.is_valid():
                        form.save()  # Save the form data to the database
                        return redirect('success')  # Redirect to success page

    except Client_Account.DoesNotExist:
        context['error_message'] = 'Client account not found.'
        return render(request, 'core/client_account_not_found.html', context)

    context = {
        'form': form,
        'header': f"Enter the following details for {client_account.ACCT_NAME}",
        'account_number': account_number,
        'client_account': client_account,
        'client_form_info': Client_Info_Form
    }

    return render(request, 'core/add_client_info.html', context)


def success_view(request):
    context = {}
    context = {'message': 'Your form has been submitted successfully!'}
    return render(request, 'core/success.html', context)

def success1(request):
    context = {}
    context = {'message': 'You have successfully generated the word document, check your outputs folder!'}
    return render(request, 'core/success1.html', context)

def generate(request):
    """View function for the generate page."""
    if request.method == 'POST':
        account_number = request.POST.get('account_number')
        try:
            # Query the database to get account data
            account = Client_Account.objects.get(ACCTNUM=account_number)
            account2 = Client_Form_Info.objects.get(ACCTNUM=account_number)  
            client_data1 = Client_Account.objects.get(ACCTNUM=account_number)
            client_data2 = Client_Form_Info.objects.get(ACCTNUM=account_number)

  # Combine data (if data found in both tables)
            if client_data1 and client_data2:
                excel_report = Excel_Report(
                    ACCTNUM=account_number,
                    ACCT_NAME=client_data1.ACCT_NAME,
                    LOAN_BALANCE=client_data1.LCYBALANCE,
                    DISBURSED_AMOUNT=client_data1.DIS_AMT,
                    DATE_OF_DISBURSEMENT=client_data1.DIS_SHDL_DATE,
                    AMOUNT_CLAIMED=client_data1.AMOUNT_CLAIMED,
                    LOAN_CYCLE=client_data1.LOAN_CYCLE,
                    DATE_OF_DEFAULT=client_data1.DATE_ARREARS_START,
                    GENDER=client_data1.GENDER,
                    AGE=client_data1.AGE,
                    SECTOR_OF_VALUE_CHAIN=client_data1.INDUSTRY_SECTOR,
                    MODE_OF_ENGAGEMENT=client_data1.MODE_OF_ENGAGEMENT,
                    REASON_FOR_DEFAULT=client_data2.REASON_FOR_DEFAULT,
                    CONTACT=client_data2.CONTACT,
                    )

    # Save data to the new table (consider using a separate function for saving)
                excel_report.save()
                messages.success(request, 'Excel report generated successfully!')

        except Client_Account.DoesNotExist:
                error_message = 'Client account not found.'
                return render(request, 'core/generate.html', {'error_message': error_message})
               

        try:
            current_directory = os.path.dirname(__file__)
    
    # Construct the full path to the template file
            templatepath = os.path.join(current_directory, 'wordtemplate', 'emplate.docx')
    
    # Load the document
            document = Document(templatepath)
                #account2 = Client_Form_Info.objects.get(ACCTNUM=account_number)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{account_name}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{account_name}}", str(account.ACCT_NAME))
                            if '{{dob}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{dob}}", str(account.CUST_DOB))
                            if '{{age}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{age}}", str(account.AGE))
                            if '{{disbursed_amount}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{disbursed_amount}}", str(account.DIS_AMT))
                            if '{{disbursement_date}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{disbursement_date}}", str(account.DIS_SHDL_DATE))
                            if '{{gender}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{gender}}", str(account.GENDER))
                            if '{{term}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{term}}", str(account.TERM))
                            if '{{claimed_amount}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{claimed_amount}}", str(account.AMOUNT_CLAIMED))
                            if '{{arrears}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{arrears}}",str(account.ARREARSDAYS))
                            if '{{business_financed}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{business_financed}}",str(account.BUSINESS_FINANCED))
                                
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{loan_purpose}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{loan_purpose}}",str(account2.LOAN_PURPOSE))
                            if '{{agroup}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{agroup}}",str(account2.GROUP))
                            if '{{app_date}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{app_date}}",str(account2.LOAN_APP_DATE))
                            if '{{cause_of_default}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{cause_of_default}}",str(account2.REASON_FOR_DEFAULT))
                            if '{{address}}' in paragraph.text:
                                paragraph.text = paragraph.text.replace("{{address}}",str(account2.ADDRESS))
                                downloads_directory = os.path.join(os.path.expanduser('~'), 'Downloads')                    
                                output_filename = f"output_{account.ACCT_NAME}.docx"
                                output_file_path = os.path.join(downloads_directory, output_filename)
                                document.save(output_file_path)
                                return redirect('success1')

                            
                        
            else:
                return HttpResponse("Failed to generate document", status=500)
        except Exception as e:
            return HttpResponse(f"An error occurred: {e}", status=500)
    else:
        return render(request, 'core/generate.html')



def excelreport(request):
  workbook = Workbook()
  worksheet = workbook.active
  worksheet.title = "Exported Data"

  # Define header row
  header_row = ["ACCTNUM", "ACCT_NAME", "LOAN_BALANCE", "DISBURSED_AMOUNT", "DATE_OF_DISBURSEMENT", "AMOUNT_CLAIMED", "DATE_OF_DEFAULT", "GENDER", "AGE", "SECTOR_OF_VALUE_CHAIN", "MODE_OF_ENGAGEMENT", "LOAN_CYCLE", "REASON_FOR_DEFAULT", "CONTACT"]  # Replace with your field names
  for col, field_name in enumerate(header_row, 1):
    worksheet.cell(row=1, column=col).value = field_name

  # Query and iterate through data
  data = Excel_Report.objects.all()
  for row, obj in enumerate(data, 2):
    for col, field in enumerate(Excel_Report._meta.fields, 1):
      value = getattr(obj, field.name)
      worksheet.cell(row=row, column=col).value = value

  # Set content type and filename
  response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
  response['Content-Disposition'] = 'attachment; filename=data.xlsx'

  workbook.save(response)

  # Display success message
  #success_message = "Excel document downloaded successfully, check your Downloads folder."

  return response

